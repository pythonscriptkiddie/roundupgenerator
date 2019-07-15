#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Thu May 23 13:45:44 2019

@author: thomassullivan
"""
import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from objects.objects import Article, Category

#class ExportedCategory:
#    def __init__(self, name, articles=[]):
#        self.name = name
#        self.articles = [article for article in articles]
#
#class ExportedArticle:
#    def __init__(self, title, url, description):
#        self.title = title
#        self.url = url
#        self.description = description

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink

def add_article(document, article):
    new_paragraph = document.add_paragraph('') #add blank paragraph that we append the text to
    add_hyperlink(new_paragraph, article.title, article.url)
    new_paragraph.add_run(' ({0} {1} {2}) '.format(article.day, article.month_text, article.year)) #blank space between the link and the description
    new_paragraph.add_run(article.description)
    
def add_category(document, category):
    category_name = document.add_paragraph(category.name)
    for article in category.articles:
        add_article(document, article)
        
def create_roundup_docx(document, roundup_title, categories):
    title = document.add_paragraph(roundup_title)
    for category in categories:
        add_category(document, category)
        
def create_complete_roundup(filename, roundup_title, categories):
    new_document = docx.Document()
    create_roundup_docx(new_document, roundup_title, categories)
    new_document.save('{0}.docx'.format(filename))
    

if __name__ == '__main__':
    
    test_articles = []
    test_categories = []
    new_document_filename = 'roundup_function_test1'

    new_roundup_title = 'OMTR Roundup Title'

    new_article = ExportedArticle(title = 'Somaliland celebrates 28th National Day in Addis',
                              url = 'https://somalilandstandard.com/somaliland-celebrates-28th-national-day-in-addis/',
                              description = 'Somaliland\'s Ambassador to Ethiopia Salan Hassan Abdilleh celebrated Somaliland\'s 28th national day.')
    
    new_article2 = ExportedArticle(title= 'BBC', url='http://www.bbc.com',
                           description = 'British Broadcasting service')
    
    test_articles.append(new_article)
    test_articles.append(new_article2)
        
    test_articles.append(new_article)
    test_articles.append(new_article2)
    
    new_category = ExportedCategory(name='Test Category 1', articles = [])
    new_category.articles.append(new_article)
    new_category.articles.append(new_article2)
    
    new_category2 = ExportedCategory(name='Test Category 2', articles = [])
    new_category2.articles.append(new_article)
    new_category2.articles.append(new_article2)
    
    test_categories.append(new_category)
    test_categories.append(new_category2)
    
    create_complete_roundup(new_document_filename, new_roundup_title, test_categories)
    
    #document = docx.Document()
    #a = document.add_paragraph('OMTR Roundup')
    #p = document.add_paragraph('')
    #for category in test_categories:
    #    add_category(document, category)
    
    #create_roundup_docx(document, new_roundup_title, test_categories)
    
    #add_hyperlink(p, new_article.title, new_article.url)
    #p.add_run(' ')
    #p.add_run(new_article.description)
    #document.save('roundup_test2.docx')